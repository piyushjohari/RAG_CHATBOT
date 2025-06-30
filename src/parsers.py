import os
import logging
from typing import Optional, Optional
from enum import auto, Enum

import pdfplumber
import pandas as pd
import tabula
import camelot
from pdf2image import convert_from_path
from paddleocr import PPStructureV3
from docx import Document

from src.utils import clean_llm_csv
from src.llm_calls import structure_table_data


class ParserType(Enum):
    PDFPLUMBER = auto()
    TABULA = auto()
    CAMELOT = auto()
    PADDLEOCR = auto()
    DOCX = auto()

    @classmethod
    def from_str(cls, name: str):
        """Get corresponding enum from string, case-insensitive."""
        name = name.strip().upper()
        for member in cls:
            if member.name == name:
                return member
        raise ValueError(f"Unknown parser type: '{name}'. Valid: {[m.name for m in cls]}")
    
class Parsers:
    """
    A class providing various PDF parsing utilities using different libraries.

    Attributes:
        folder_path (str): The default folder containing PDF files.
        output_dir (str): The directory where parsed tables/data are saved.
        poppler_path (Optional[str]): Path to Poppler if required by pdf2image (Windows).
    """

    def __init__(self, 
                 folder_path: str = "static/insurace_data", 
                 output_dir: str = "static/parsed_data/pdf_table",
                 poppler_path: Optional[str] = None):
        self.folder_path = folder_path
        self.output_dir = output_dir
        self.poppler_path = poppler_path

        os.makedirs(self.output_dir, exist_ok=True)
        logging.basicConfig(level=logging.INFO, 
                            format='%(asctime)s - %(levelname)s - %(message)s')
        

    def parse(self, file_path: str, parser: str):
        """
        Dispatches to the selected parser according to the parser type string.
        """
        try:
            parser_enum = ParserType.from_str(parser)
        except ValueError as ve:
            logging.error(str(ve))
            raise

        if parser_enum == ParserType.PDFPLUMBER:
            self.parse_pdfplumber(file_path)
        elif parser_enum == ParserType.TABULA:
            self.parse_tabula(file_path)
        elif parser_enum == ParserType.CAMELOT:
            self.parse_camelot(file_path)
        elif parser_enum == ParserType.PADDLEOCR:
            self.parse_paddleocr(file_path)
        elif parser_enum == ParserType.DOCX:
            self.parse_docx(file_path)

    def parse_pdfplumber(self, file_path: str):
        """
        Parses PDF using pdfplumber.
        Extracts text and tables, saves results to output_dir.
        """
        try:
            output_dir_txt = 'static/parsed_data/pdf_text'
            output_dir_csv = 'static/parsed_data/pdf_table'
            file_name = os.path.basename(file_path)
            full_text = ""
            with pdfplumber.open(file_path) as pdf:
                for page_number, page in enumerate(pdf.pages, start=1):
                    text = page.extract_text()

                    if text and text.strip():
                        full_text = full_text + f'\n--- Page: {page_number} ---\n{text}'
                        tables = page.extract_tables()

                        if tables:
                            for t_idx, table in enumerate(tables, start=1):
                                df = pd.DataFrame(table)
                                structured_data = structure_table_data(df, text, file_name)
                                cleaned_data = clean_llm_csv(structured_data)
                                out_file = os.path.join(
                                    output_dir_csv, f"{file_name.replace('.pdf', '')}_page{page_number}_table{t_idx}.csv")
                                os.makedirs(os.path.dirname(out_file), exist_ok=True)
                                with open(out_file, "w", encoding="utf-8") as f:
                                    f.write(cleaned_data)
                                logging.info(f"Saved structured table to {out_file}")
                        else:
                            logging.info(f"No tables found on page {page_number} of {file_name}.")
                    else:
                        logging.info(f"No extractable text on page {page_number} of {file_name}.")
            output_file_path = os.path.join(output_dir_txt, file_name.replace('.pdf', '.txt'))
            os.makedirs(os.path.dirname(output_file_path), exist_ok=True)
            with open(output_file_path, 'w', encoding='utf-8') as f:
                f.write(full_text)
            print(f"Text written to {output_file_path}")
        except Exception as e:
            logging.error(f"Failed on {file_name}: {e}")
        logging.info("Completed pdfplumber parsing.")


    def parse_docx(self, file_path: str) -> str:
        """
        Reads a .docx file and returns the full text content as a single string.
        
        Args:
            file_path: Path to the .docx file.
        
        Returns:
            All text content, separated by newlines.
        """
        output_dir = 'static/parsed_data/docx'
        file_name = os.path.basename(file_path)
        doc = Document(file_path)
        # Extract from paragraphs
        full_text = [para.text for para in doc.paragraphs]

        # Extract text from tables as well (optional)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        full_text.append(cell_text)

        # Join all extracted text together
        content = '\n'.join(line for line in full_text if line.strip())
        output_file_path = os.path.join(output_dir, file_name.replace('.docx', '.txt'))
        os.makedirs(os.path.dirname(output_file_path), exist_ok=True)
        with open(output_file_path, 'w', encoding='utf-8') as f:
            f.write(content)
        print(f"Text written to {output_file_path}")


    def parse_camelot(self, file_path: str):
        """
        Parses a PDF file using Camelot. Extracts tables and saves as CSV.
        """

        logging.info(f"Processing (camelot): {file_path}")
        try:
            tables = camelot.read_pdf(file_path, pages='all', flavor='stream')
            file_name = os.path.basename(file_path)
            for i, table in enumerate(tables):
                out_file = os.path.join(
                    self.output_dir, f"{file_name}_camelot_{i+1}.csv")
                os.makedirs(os.path.dirname(out_file), exist_ok=True)
                table.to_csv(out_file)
                logging.info(f"Saved Camelot table {i+1} to {out_file}")
        except Exception as e:
            logging.error(f"Camelot failed on {file_path}: {e}")

    def parse_tabula(self, file_path: str):
        """
        Parses all PDF files in the folder using Tabula, 
        extracting all tables from all pages and saving optionally.
        """
        # for file_name in os.listdir(self.folder_path):
        if file_path.lower().endswith('.pdf'):
            file_name = os.path.basename(file_path)
            logging.info(f"Processing (tabula): {file_path}")
            try:
                tables = tabula.read_pdf(
                    file_path,
                    pages='all',
                    multiple_tables=True,
                    lattice=True
                )
                for idx, df in enumerate(tables, 1):
                    out_file = os.path.join(
                        self.output_dir, f"{file_name}_tabula_{idx}.csv")
                    df.to_csv(out_file, index=False)
                    logging.info(f"Saved table {idx} to {out_file}")
            except Exception as e:
                logging.error(f"Tabula failed on {file_name}: {e}")
        logging.info("Completed Tabula parsing.")


    def parse_paddleocr(self, file_path: str):
        """
        Uses PaddleOCR to perform table structure detection in scanned PDFs.
        Converts PDF pages to images and applies OCR.
        """

        logging.info(f"Processing (paddleocr): {file_path}")
        try:
            # Convert PDF pages to images
            file_name = os.path.basename(file_path)
            pages = convert_from_path(file_path, poppler_path=self.poppler_path)
            ocr = PPStructureV3()
            for idx, page_img in enumerate(pages):
                img_filename = os.path.join(self.output_dir, f"{os.path.splitext(file_name)[0]}_paddle_page{idx+1}.jpg")
                page_img.save(img_filename)
                result = ocr(img_filename)
                ocr.write_result(result, output=self.output_dir, img_name=f"{os.path.splitext(file_name)[0]}_paddle_page{idx+1}")
                logging.info(f"PaddleOCR results written for page {idx+1}")
        except Exception as e:
            logging.error(f"PaddleOCR failed on {file_name}: {e}")
